import streamlit as st
import os
import pdfkit
from docx import Document
from google import genai
import tempfile
from dotenv import load_dotenv  # <--- NEW IMPORT

# Load variables from .env file immediately
load_dotenv()
# --- CONFIGURATION ---
# Set your API Key here or in environment variables
# If using environment variable: api_key = os.environ.get("GEMINI_API_KEY")
API_KEY = os.environ.get("GEMINI_API_KEY")

# Configure PDFKit path (Update this if needed, or set to None if in PATH)
WKHTMLTOPDF_PATH = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"

# Initialize Gemini
client = genai.Client(api_key=API_KEY)

# --- HELPER FUNCTIONS ---
def convert_to_pdf(content):
    """Converts HTML string to PDF bytes."""
    try:
        config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH) if WKHTMLTOPDF_PATH else None
        html_wrapper = f"""
        <html>
          <head><meta charset="UTF-8"></head>
          <body>{content}</body>
        </html>
        """
        # Return PDF as bytes
        return pdfkit.from_string(html_wrapper, False, configuration=config)
    except Exception as e:
        st.error(f"PDF Error: {e}")
        return None

def convert_to_word(content):
    """Converts text to Word Doc bytes."""
    try:
        doc = Document()
        doc.add_heading('Gemini Processed Report', 0)
        doc.add_paragraph(content)
        
        # Save to a temporary file to read bytes back
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
            
        with open(tmp_path, "rb") as f:
            docx_bytes = f.read()
        
        os.unlink(tmp_path) # Clean up
        return docx_bytes
    except Exception as e:
        st.error(f"Word Error: {e}")
        return None

# --- STREAMLIT UI ---
st.set_page_config(page_title="AI Document Processor", page_icon="📄")

st.title("📄 AI HTML By chandrika Agent")
st.write("Upload raw HTML files, define your requirements, and let Gemini process them.")

# 1. SIDEBAR: Configuration
with st.sidebar:
    st.header("⚙️ Settings")
    output_format = st.radio("Output Format:", ["PDF", "Word Document (.docx)"])
    
    st.divider()
    
    st.subheader("Processing Mode")
    mode = st.selectbox(
        "Choose Action:",
        ["Summarize Content", "Clean & Format", "Extract Key Data", "Custom Instruction"]
    )
    
    custom_instruction = ""
    if mode == "Custom Instruction":
        custom_instruction = st.text_area("Enter custom prompt:", "E.g., Translate to Spanish...")

# 2. FILE UPLOAD
uploaded_file = st.file_uploader("Upload an HTML file", type=["html", "htm"])

if uploaded_file is not None:
    # Read file
    string_data = uploaded_file.getvalue().decode("utf-8")
    
    with st.expander("Preview Raw HTML Content"):
        st.code(string_data[:1000] + "...", language='html')

    # 3. PROCESS BUTTON
    if st.button("🚀 Process File"):
        with st.spinner("Gemini is thinking..."):
            
            # Construct Prompt based on User Selection
            base_prompt = ""
            if mode == "Summarize Content":
                base_prompt = "Summarize the following content into a professional report."
            elif mode == "Clean & Format":
                base_prompt = "Clean up this content and format it nicely with headers and bullet points."
            elif mode == "Extract Key Data":
                base_prompt = "Extract key metrics, dates, and names from this content."
            else:
                base_prompt = custom_instruction

            full_prompt = (
                f"{base_prompt}\n"
                "IMPORTANT: Output the result as clean HTML (using tags like <h3>, <p>, <ul>) "
                "so it can be converted to a document. Do not use Markdown blocks."
                f"\n\nCONTENT:\n{string_data}"
            )

            try:
                # Call Gemini
                response = client.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=full_prompt
                )
                processed_text = response.text
                
                st.success("Processing Complete!")
                
                # Show Result Preview
                st.subheader("📝 Result Preview")
                st.markdown(processed_text, unsafe_allow_html=True)

                # 4. DOWNLOAD BUTTON
                file_name_root = os.path.splitext(uploaded_file.name)[0]
                
                if output_format == "PDF":
                    pdf_data = convert_to_pdf(processed_text)
                    if pdf_data:
                        st.download_button(
                            label="📥 Download PDF",
                            data=pdf_data,
                            file_name=f"{file_name_root}_processed.pdf",
                            mime="application/pdf"
                        )
                
                elif output_format == "Word Document (.docx)":
                    docx_data = convert_to_word(processed_text)
                    if docx_data:
                        st.download_button(
                            label="📥 Download Word Doc",
                            data=docx_data,
                            file_name=f"{file_name_root}_processed.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            except Exception as e:
                st.error(f"An error occurred: {e}")